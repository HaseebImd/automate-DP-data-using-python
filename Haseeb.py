import random
import docx
def matrix_chain_multiplication(p):
    n = len(p) - 1
    m = [[0 for i in range(n + 1)] for j in range(n + 1)]
    s = [[0 for i in range(n + 1)] for j in range(n + 1)]
    for l in range(2, n + 1):
        for i in range(1, n - l + 2):
            j = i + l - 1
            m[i][j] = float("inf")
            for k in range(i, j):
                q = m[i][k] + m[k + 1][j] + (p[i - 1] * p[k] * p[j])
                if q < m[i][j]:
                    m[i][j] = q
                    s[i][j] = k
    return m, s

def print_optimal_parens(s, i, j):
    sequence = ""
    if i == j:
        sequence += "A" + str(i)
    else:
        sequence += "("
        sequence += print_optimal_parens(s, i, s[i][j])
        sequence += print_optimal_parens(s, s[i][j] + 1, j)
        sequence += ")"
    return sequence

def Calculate_P_Matrices(minimum_P_matrices,min_P_length,max_p_length,min_P_element_value,max_P_element_value):
    doc = docx.Document()
    for Matrix in range(minimum_P_matrices):
        Mat=[]
        minimum_P_length=random.randint(min_P_length, max_p_length)
        for ActualMatrix in range(minimum_P_length):
            matrix_index_value=random.randint(min_P_element_value, max_P_element_value)
            Mat.append(matrix_index_value)
        m, s = matrix_chain_multiplication(Mat)
        d=print_optimal_parens(s,1,len(Mat)-1)
        doc.add_paragraph(str(Mat))
        doc.add_paragraph('Minimum Multiplications: '+str(m[1][len(Mat)-1]))
        doc.add_paragraph('Optimal Sequence : \n'+str(d))
        doc.add_paragraph('\n')
    doc.save('Mahnor.docx')




# How many P matrices you want to generate
minimum_P_matrices=40

# Provide minimum and maximum value b/w this any random value will be generated
# and it would be the length of that P array
min_P_length=10
max_p_length=30

# Provide minimum and maximum value b/w this any random value will be generated
# and it would be the every element of array

min_P_element_value=15
max_P_element_value=35

Calculate_P_Matrices(minimum_P_matrices,min_P_length,max_p_length,min_P_element_value,max_P_element_value)



